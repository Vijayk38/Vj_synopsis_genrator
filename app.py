import os
import re
import io 
import streamlit as st 

# --- CORRECTED IMPORTS ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import google.generativeai as genai

# --- Gemini Generation Function ---

@st.cache_resource
def get_gemini_model():
    """Initializes and returns the Gemini model."""
    try:
        # Access the API key from Streamlit's secrets
        if "GOOGLE_API_KEY" not in st.secrets:
            st.error("Error: GOOGLE_API_KEY not found in Streamlit secrets.")
            st.stop()
            
        api_key = st.secrets["GOOGLE_API_KEY"]
        genai.configure(api_key=api_key)
        # UPDATED: Use a valid model name (gemini-1.5-flash)
        return genai.GenerativeModel("gemini-1.5-flash")
    except Exception as e:
        st.error(f"Error configuring Gemini: {e}")
        st.stop()


def generate_report_with_gemini(topic, model):
    """Generates the report content using the Gemini model."""
    
    prompt = f"""
    Generate a **comprehensive, detailed, and professionally structured project report** for the topic: {topic}. 
    
    CRITICAL TONE AND STYLE INSTRUCTIONS:
    1. Write in an **active voice**, demonstrating critical analysis and expert understanding.
    2. Use sophisticated academic language and natural transitions.
    3. The report must be very long and detailed (aim for 3000+ words).
    
    Structure the report with these exact headings followed by a colon:
    
    1. Introduction:
    1.1 Background and Context:
    1.2 Problem Statement:
    
    2. Literature Review:
    2.1 Current State of Research:
    2.2 Identification of Gaps:
    
    3. Project Objectives:
    3.1 Specific Aims:
    3.2 Deliverables:
    
    4. Detailed Methodology and System Design:
    4.1 Proposed Architecture:
    4.2 Data Collection and Analysis Methods:
    
    5. Implementation Details and Results Analysis: 
    5.1 Execution Plan:
    5.2 Expected Outcomes and Validation:
    
    6. Conclusion and Future Work: 
    6.1 Summary of Findings:
    6.2 Future Scope and Recommendations:

    Format: Numbered section headers followed by a colon. 
    Do NOT use Markdown characters like # or *.
    """
    
    try:
        response = model.generate_content(prompt, generation_config={"temperature": 0.7}) 
        if not response.text:
            return None, "API returned an empty response. Try a different topic."
        return response.text, None
    except Exception as e:
        return None, f"Report Generation Error: {str(e)}"

# --- Word Conversion Function ---

def text_to_word_buffer(text_content, topic):
    """Parses text and formats it into a Word document buffer."""
    try:
        document = Document()
        
        # Margins
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Title
        topic_heading = document.add_heading(topic.upper(), level=0)
        topic_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph()
        
        lines = text_content.strip().split('\n')
        
        for line in lines:
            clean_line = line.strip()
            if not clean_line:
                continue
            
            # Detect Headers (e.g., "1. Introduction:")
            header_match = re.match(r'^(\d+(\.\d+)*)\s*(.*?):$', clean_line)

            if header_match:
                numbering = header_match.group(1) 
                title = header_match.group(3).strip().upper()

                if numbering.count('.') == 0:
                    h = document.add_heading(f"{numbering} {title}", level=1)
                else:
                    h = document.add_heading(f"{numbering} {title}", level=2)
            else:
                # Normal paragraph text
                p = document.add_paragraph(clean_line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                fmt = p.paragraph_format
                fmt.line_spacing = 1.5 
        
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0) 
        
        return doc_io, None
    except Exception as e:
        return None, f"Word Conversion Error: {str(e)}"

# --- Streamlit Main App ---

def main():
    st.set_page_config(page_title="Gemini Project Report Gen", layout="centered")

    st.title("📄 AI Project Report Generator")
    st.write("Generates professional 3000+ word reports in DOCX format.")
    st.markdown("---")

    topic = st.text_input("Enter Project Topic:", placeholder="e.g. Smart Irrigation System using IoT")

    if st.button("Generate & Download Report", type="primary"):
        if not topic:
            st.error("Please enter a topic.")
            return

        model = get_gemini_model()
        
        with st.spinner("Generating detailed content... this may take a minute."):
            report_text, error = generate_report_with_gemini(topic, model)

        if error:
            st.error(error)
            return

        with st.spinner("Converting to Word..."):
            doc_buffer, word_error = text_to_word_buffer(report_text, topic)

        if word_error:
            st.error(word_error)
            return

        st.success("✅ Done!")
        
        st.download_button(
            label="Download DOCX Report",
            data=doc_buffer,
            file_name=f"{topic.replace(' ', '_')}_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
